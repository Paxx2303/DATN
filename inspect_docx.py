# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

DOC = r"E:\test\DATN\YOLOv11_BaoCao_TiengViet_v3_modified.docx"

doc = Document(DOC)

print(f"=== PARAGRAPHS ({len(doc.paragraphs)}) ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    style = p.style.name if p.style else "None"
    has_img = any(
        child.tag.split('}')[-1] in ('drawing', 'pict')
        for child in p._p.iter()
    )
    marker = " [IMG]" if has_img else ""
    has_pb = any(
        br.get(qn("w:type"), "") == "page"
        for br in p._p.findall(".//" + qn("w:br"))
    )
    pb_marker = " [PAGEBREAK]" if has_pb else ""
    if t or has_img or has_pb:
        print(f"  [{i:4d}] style={style!r:30s}{pb_marker}{marker} | {t[:100]}")

print(f"\n=== TABLES ({len(doc.tables)}) ===")
for i, tbl in enumerate(doc.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns) if tbl.columns else 0
    first_cell = tbl.rows[0].cells[0].text[:60] if rows > 0 else ""
    print(f"  [Table {i}] {rows}r x {cols}c | first_cell: {first_cell!r}")

print("\n=== STYLES (paragraph styles in use) ===")
used_styles = set()
for p in doc.paragraphs:
    if p.style:
        used_styles.add(p.style.name)
for s in sorted(used_styles):
    print(f"  {s!r}")
