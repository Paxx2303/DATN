import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def insert_toc_field(p, toc_instruction):
    p.clear()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = toc_instruction
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

doc = docx.Document(r'C:\Using\DATN\YOLOv11_BaoCao_Updated.docx')
styles = doc.styles

try:
    style_bang = styles.add_style('CaptionBang', WD_STYLE_TYPE.PARAGRAPH)
    style_bang.base_style = styles['Normal']
except Exception:
    pass

try:
    style_hinh = styles.add_style('CaptionHinh', WD_STYLE_TYPE.PARAGRAPH)
    style_hinh.base_style = styles['Normal']
except Exception:
    pass

state = 'normal'
toc_bang_p = None
toc_hinh_p = None

paragraphs_to_remove = []

for p in doc.paragraphs:
    text = p.text.strip()
    
    if text == 'DANH MỤC BẢNG BIỂU':
        state = 'toc_bang'
        continue
    elif text == 'DANH MỤC HÌNH VẼ' or text == 'DANH MỤC HÌNH ẢNH':
        state = 'toc_hinh'
        continue
    elif text == 'DANH MỤC TỪ VIẾT TẮT' or text == 'MỞ ĐẦU':
        state = 'normal'
    
    if state == 'toc_bang':
        if text.startswith('Bảng ') or not text:
            paragraphs_to_remove.append(p)
            if not toc_bang_p:
                toc_bang_p = p # use the first one to inject TOC field
        else:
            state = 'normal'
    elif state == 'toc_hinh':
        if text.startswith('Hình ') or not text:
            paragraphs_to_remove.append(p)
            if not toc_hinh_p:
                toc_hinh_p = p
        else:
            state = 'normal'
            
    if state == 'normal' and text:
        if text.startswith('Bảng '):
            p.style = 'CaptionBang'
        elif text.startswith('Hình '):
            p.style = 'CaptionHinh'

if toc_bang_p:
    insert_toc_field(toc_bang_p, ' TOC \\h \\z \\t "CaptionBang,1" ')
if toc_hinh_p:
    insert_toc_field(toc_hinh_p, ' TOC \\h \\z \\t "CaptionHinh,1" ')

for p in paragraphs_to_remove:
    if p != toc_bang_p and p != toc_hinh_p:
        p_element = p._element
        p_element.getparent().remove(p_element)

# Enable updateFields
settings = doc.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

doc.save(r'C:\Using\DATN\YOLOv11_BaoCao_Updated.docx')
print('Done!')
