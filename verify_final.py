# -*- coding: utf-8 -*-
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

doc = Document(r"E:\test\DATN\YOLOv11_BaoCao_TiengViet_Final.docx")

print("=== HEADING STYLES ===")
for i, p in enumerate(doc.paragraphs):
    sname = p.style.name if p.style else ''
    t = p.text.strip()
    if 'Heading' in sname and t:
        print(f"  [{i:4d}] {sname} | {t[:80]}")

print("\n=== IMAGES ===")
img_count = 0
for i, p in enumerate(doc.paragraphs):
    has_img = any(c.tag.split('}')[-1] in ('drawing','pict') for c in p._p.iter())
    if has_img:
        img_count += 1
        print(f"  [{i:4d}] [IMG] next: {doc.paragraphs[i+1].text[:60] if i+1<len(doc.paragraphs) else ''}")
print(f"Total images: {img_count}")

print("\n=== REMAINING PLACEHOLDERS [Hình ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("[Hình"):
        print(f"  [{i:4d}] {p.text[:80]}")

print("\n=== TOC FIELD ===")
for i, p in enumerate(doc.paragraphs):
    instr_els = p._p.findall(".//" + qn("w:instrText"))
    for el in instr_els:
        if el.text and 'TOC' in el.text:
            print(f"  [{i:4d}] Found TOC field: {el.text[:60]}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Check for remaining duplicate captions
print("\n=== DUPLICATE CAPTION CHECK ===")
dups = 0
paras = doc.paragraphs
for i in range(len(paras)-1):
    t1 = paras[i].text.strip(); t2 = paras[i+1].text.strip()
    m1 = re.match(r'^Hình (\d+\.\d+)', t1); m2 = re.match(r'^Hình (\d+\.\d+)', t2)
    if m1 and m2 and m1.group(1) == m2.group(1):
        print(f"  DUPLICATE at [{i}]: {t1[:60]}")
        dups += 1
print(f"  {dups} duplicate pairs remaining")
