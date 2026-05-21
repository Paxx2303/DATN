# -*- coding: utf-8 -*-
"""
format_docx_v2.py
- Bắt đầu từ file gốc (không có ảnh)
- Chỉ chèn ảnh khi nội dung thực sự khớp
- Xóa các trang trắng / đoạn trống thừa
- Đổi toàn bộ font → Times New Roman 13pt
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_IN  = r"e:\test\DATN\YOLOv11_BaoCao_TiengViet.docx"   # file gốc
DOC_OUT = r"e:\test\DATN\YOLOv11_BaoCao_TiengViet_v3.docx"
IMG_DIR = r"e:\test\DATN\report_imgs"

FONT_NAME = "Times New Roman"
FONT_SIZE = 13

# Chỉ giữ ảnh thực sự khớp với nội dung
IMAGE_MAP = {
    "Hình 2.3": ("chart_fisheye_distortion.png", 13,
                 "Hình 2.3. Kết quả áp dụng hàm to_fisheye() trên ảnh giao thông (strength=0,5)"),
    "Hình 2.5": ("chart_yolo_arch.png",          14,
                 "Hình 2.5. Kiến trúc khối C3k2 (Cross Stage Partial with kernel size 2)"),
    "Hình 3.1": ("chart_system_arch.png",        15,
                 "Hình 3.1. Sơ đồ kiến trúc hệ thống tổng thể"),
    "Hình 4.4": ("chart_loss_curves.png",        13,
                 "Hình 4.4. Đường cong training loss và validation loss theo epoch"),
    "Hình 4.5": ("chart_map_curves.png",         13,
                 "Hình 4.5. Đường cong mAP@0.5 và mAP@0.5:0.95 theo epoch"),
    "Hình 4.6": ("chart_confusion_matrix.png",   12,
                 "Hình 4.6. Confusion matrix trên tập kiểm thử (normalized)"),
    "Hình 4.7": ("chart_model_compare.png",      13,
                 "Hình 4.7. So sánh hiệu năng các biến thể YOLOv11 trên tập kiểm thử"),
}


# ─── Font helpers ────────────────────────────────────────────────────────────

def _set_rFonts(rPr_el):
    fonts_el = rPr_el.find(qn("w:rFonts"))
    if fonts_el is None:
        fonts_el = OxmlElement("w:rFonts")
        rPr_el.insert(0, fonts_el)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts_el.set(qn(attr), FONT_NAME)


def _set_sz(rPr_el, pt):
    for tag in ("w:sz", "w:szCs"):
        el = rPr_el.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr_el.append(el)
        el.set(qn("w:val"), str(pt * 2))


def set_para_font(para):
    # Paragraph mark (pPr/rPr)
    pPr = para._p.get_or_add_pPr()
    rPr_pm = pPr.find(qn("w:rPr"))
    if rPr_pm is None:
        rPr_pm = OxmlElement("w:rPr")
        pPr.append(rPr_pm)
    _set_rFonts(rPr_pm)
    _set_sz(rPr_pm, FONT_SIZE)

    # Individual runs
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        rPr = run._r.get_or_add_rPr()
        _set_rFonts(rPr)


def format_tables(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_para_font(para)


# ─── Image insertion ─────────────────────────────────────────────────────────

def insert_image_after(doc, anchor_p_el, img_path, width_cm, caption_text):
    """Chèn paragraph ảnh + caption sau anchor_p_el."""
    from docx.text.paragraph import Paragraph

    # Paragraph ảnh (căn giữa)
    img_p_el = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    pPr.append(jc)
    img_p_el.append(pPr)
    anchor_p_el.addnext(img_p_el)

    p_wrapper = Paragraph(img_p_el, doc)
    run = p_wrapper.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    # Paragraph caption (italic, 12pt, căn giữa)
    cap_p_el = OxmlElement("w:p")
    cap_pPr = OxmlElement("w:pPr")
    cap_jc = OxmlElement("w:jc"); cap_jc.set(qn("w:val"), "center")
    cap_pPr.append(cap_jc)
    cap_p_el.append(cap_pPr)

    run_el = OxmlElement("w:r")
    run_rPr = OxmlElement("w:rPr")
    fonts_el = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts_el.set(qn(attr), FONT_NAME)
    run_rPr.append(fonts_el)
    sz  = OxmlElement("w:sz");   sz.set(qn("w:val"),   "24")
    szC = OxmlElement("w:szCs"); szC.set(qn("w:val"),  "24")
    i   = OxmlElement("w:i")
    run_rPr.extend([sz, szC, i])
    run_el.append(run_rPr)

    t_el = OxmlElement("w:t")
    t_el.text = caption_text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_el.append(t_el)
    cap_p_el.append(run_el)
    img_p_el.addnext(cap_p_el)


# ─── Blank-page / empty-para cleanup ─────────────────────────────────────────

def _is_empty_para(p_el):
    """True nếu paragraph không có text thực và không chứa ảnh."""
    for child in p_el.iter():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("drawing", "pict"):
            return False
    text = "".join(t.text or "" for t in p_el.findall(".//" + qn("w:t")))
    return text.strip() == ""


def _has_page_break(p_el):
    """True nếu paragraph chứa page break hoặc là section break tạo trang mới."""
    for br in p_el.findall(".//" + qn("w:br")):
        if br.get(qn("w:type"), "") == "page":
            return True
    # Section break (continuous là OK, nhưng nextPage/oddPage/evenPage tạo trang mới)
    sectPr = p_el.find(".//" + qn("w:sectPr"))
    if sectPr is not None:
        pbSect = sectPr.find(qn("w:type"))
        if pbSect is None:
            return True  # default = nextPage
        val = pbSect.get(qn("w:val"), "nextPage")
        if val in ("nextPage", "oddPage", "evenPage"):
            return True
    return False


def remove_blank_pages(doc):
    """
    Xóa các trang trắng dư thừa:
    1. Giảm chuỗi >2 đoạn trống liên tiếp xuống còn tối đa 1.
    2. Xóa page-break paragraph nếu paragraph ngay sau đó trống (trang trắng).
    """
    body = doc.element.body
    children = list(body)

    removed = 0

    # Bước 1: Thu gọn chuỗi đoạn trống liên tiếp
    consecutive_empty = 0
    to_remove = []
    for el in children:
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p" and _is_empty_para(el):
            consecutive_empty += 1
            if consecutive_empty > 1:
                to_remove.append(el)
        else:
            consecutive_empty = 0
    for el in to_remove:
        body.remove(el)
        removed += 1

    # Bước 2: Xóa page-break paragraph tạo trang trắng
    # (page-break theo sau bởi đoạn trống hoặc page-break khác)
    children = list(body)
    to_remove2 = []
    for i, el in enumerate(children[:-1]):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag != "p":
            continue
        if not _has_page_break(el):
            continue
        # Kiểm tra có phải page-break thuần không (không có text thực)
        text = "".join(t.text or "" for t in el.findall(".//" + qn("w:t")))
        if text.strip():
            continue
        # Nếu trước nó là trang trắng / page-break khác → xóa
        next_el = children[i + 1]
        next_tag = next_el.tag.split("}")[-1] if "}" in next_el.tag else next_el.tag
        if next_tag == "p" and (_is_empty_para(next_el) or _has_page_break(next_el)):
            to_remove2.append(el)
        # Nếu nó là đoạn trống với page-break duy nhất → xóa
        elif _is_empty_para(el):
            to_remove2.append(el)

    for el in to_remove2:
        if el in list(body):
            body.remove(el)
            removed += 1

    return removed


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    print("Mo tai lieu goc...")
    doc = Document(DOC_IN)

    # 1. Đặt font mặc định Normal style
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = Pt(FONT_SIZE)

    # 2. Font từng paragraph
    print("Doi font paragraphs...")
    for para in doc.paragraphs:
        set_para_font(para)

    # 3. Font bảng
    print("Doi font bang...")
    format_tables(doc)

    # 4. Chèn ảnh khớp nội dung
    print("Chen hinh anh...")
    paras = list(doc.paragraphs)
    inserted = 0

    for idx, para in enumerate(paras):
        text = para.text.strip()
        if not (text.startswith("[H") or text.startswith("[h")):
            continue

        matched_key = None
        for key in IMAGE_MAP:
            if key in text:
                matched_key = key
                break
        if matched_key is None:
            continue

        img_file, width_cm, caption = IMAGE_MAP[matched_key]
        img_path = os.path.join(IMG_DIR, img_file)
        if not os.path.exists(img_path):
            print(f"  [SKIP] Khong tim thay anh: {img_path}")
            continue

        print(f"  [OK] {matched_key} -> {img_file}")

        # Xóa text placeholder, giữ element
        p_el = para._p
        for child in list(p_el):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r", "bookmarkStart", "bookmarkEnd", "proofErr"):
                p_el.remove(child)
        for t in p_el.findall(".//" + qn("w:t")):
            t.getparent().remove(t)

        try:
            insert_image_after(doc, p_el, img_path, width_cm, caption)
            inserted += 1
        except Exception as e:
            print(f"  [ERR] {matched_key}: {e}")

    print(f"Chen {inserted} hinh anh.")

    # 5. Xóa trang trắng
    print("Xoa trang trang...")
    removed = remove_blank_pages(doc)
    print(f"Da xoa {removed} phan tu trang / trong.")

    # 6. Lưu
    print(f"Luu -> {DOC_OUT}")
    doc.save(DOC_OUT)
    print("Hoan thanh!")


if __name__ == "__main__":
    main()
