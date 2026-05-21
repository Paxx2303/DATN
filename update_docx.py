"""
Script to update YOLOv11_Report_Final (1).docx with v5 notebook parameters.
Key changes based on graduate-project-v5-resume-v3 notebook:
  - Phase 2: 2x Tesla T4, 80 additional epochs (130 total), lr=0.001, lrf=0.002, warmup=3, patience=40, batch=32
  - Final results: mAP@0.5=0.616, mAP@0.5:0.95=0.368, P=0.654, R=0.582, ~181 FPS
  - Per-class: Car(0.776), Bus(0.565), Truck(0.515), Pedestrian(0.339), Motorbike(0.887)
"""

import copy
import docx
from docx import Document
from docx.shared import Pt
import re

INPUT_PATH  = r"c:\Using\NCKH\YOLOv11_Report_Final (1).docx"
OUTPUT_PATH = r"c:\Using\NCKH\YOLOv11_Report_Final_v5.docx"

doc = Document(INPUT_PATH)

# ─────────────────────────────────────────────
# Helper: replace text in a paragraph preserving runs
# ─────────────────────────────────────────────
def replace_para_text(para, old, new):
    """Replace a substring across the full paragraph text, preserving run formatting."""
    full = para.text
    if old not in full:
        return False
    # Simplest safe approach: rewrite first run, clear rest
    new_full = full.replace(old, new)
    # Store formatting from first run
    if para.runs:
        fmt = para.runs[0]
        bold = fmt.bold
        size = fmt.font.size
        color = fmt.font.color.rgb if fmt.font.color and fmt.font.color.type else None
    else:
        bold, size, color = None, None, None

    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
        para.runs[0].bold = bold
        if size:
            para.runs[0].font.size = size
    return True


def replace_in_all_paras(old, new):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            if replace_para_text(p, old, new):
                count += 1
    return count


# ─────────────────────────────────────────────
# 1. Update Table 3 — Training Hyperparameters
#    (was: P100, 50 epochs, lr=0.0005, warmup=5)
#    Now represent the FULL 2-phase training
# ─────────────────────────────────────────────
table3 = doc.tables[3]   # 0-indexed: Hyperparameter | Value | Notes

updates_t3 = {
    "Epochs": ("130 (50 + 80)", "Phase 1: 50 ep (P100) → Phase 2: 80 ep (2×T4)"),
    "Learning Rate": ("0.001", "Phase 2 AdamW lr0 (Phase 1: 0.0005)"),
    "Batch Size": ("32 effective", "2×16 across 2× Tesla T4 GPUs"),
    "Optimizer": ("AdamW", "weight_decay = 0.0005, lrf = 0.002"),
    "Warmup": ("3 epochs", "Phase 2 warmup (Phase 1: 5 epochs)"),
    "GPU": ("Tesla P100 + 2× Tesla T4", "Phase 1: 1×P100 | Phase 2: 2×T4 (15 GB each)"),
}

for row in table3.rows:
    cells = row.cells
    if len(cells) < 2:
        continue
    key = cells[0].text.strip()
    if key in updates_t3:
        val, note = updates_t3[key]
        # Update Value cell
        for run in cells[1].paragraphs[0].runs:
            run.text = ""
        if cells[1].paragraphs[0].runs:
            cells[1].paragraphs[0].runs[0].text = val
        else:
            cells[1].paragraphs[0].add_run(val)
        # Update Notes cell if exists
        if len(cells) >= 3:
            for run in cells[2].paragraphs[0].runs:
                run.text = ""
            if cells[2].paragraphs[0].runs:
                cells[2].paragraphs[0].runs[0].text = note
            else:
                cells[2].paragraphs[0].add_run(note)

print("✓ Table 3 (Hyperparameters) updated")


# ─────────────────────────────────────────────
# 2. Update Table 4 — Per-class performance
#    Old: Car/Bus/Bike/Truck/Pedestrian (50-ep results, mAP=0.419)
#    New: Car/Bus/Truck/Pedestrian/Motorbike (v5 results, mAP=0.616)
# ─────────────────────────────────────────────
table4 = doc.tables[4]

# New per-class data from v5 best.pt validation
new_perf = {
    # class -> (Precision, Recall, mAP@0.5, F1, Note)
    "Car":        ("0.718", "0.719", "0.776", "0.718", "Best AP class"),
    "Bus":        ("0.623", "0.569", "0.565", "0.595", ""),
    "Bike":       ("0.805", "0.858", "0.887", "0.831", "Motorbike — best Recall"),
    "Truck":      ("0.519", "0.486", "0.515", "0.502", ""),
    "Pedestrian": ("0.603", "0.276", "0.339", "0.379", "Low recall — periphery distortion"),
    "Overall (all)": ("0.654", "0.582", "0.616", "0.616", ""),
}

for row in table4.rows:
    cells = row.cells
    if len(cells) < 4:
        continue
    cls = cells[0].text.strip()
    if cls in new_perf:
        vals = new_perf[cls]
        col_data = [cls] + list(vals)
        for ci, val in enumerate(col_data):
            if ci >= len(cells):
                break
            para = cells[ci].paragraphs[0]
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = val
            else:
                para.add_run(val)

print("✓ Table 4 (Per-class performance) updated")


# ─────────────────────────────────────────────
# 3. Update Table 5 — Comparison table
#    YOLOv11-L: 0.419 → 0.616, FPS: 41 → 181
# ─────────────────────────────────────────────
table5 = doc.tables[5]

for row in table5.rows:
    cells = row.cells
    if len(cells) < 2:
        continue
    if "YOLOv11" in cells[0].text or "proposed" in cells[0].text:
        new_vals = {1: "0.616", 2: "25.3", 3: "181", 4: "86.6"}
        for ci, val in new_vals.items():
            if ci < len(cells):
                para = cells[ci].paragraphs[0]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = val
                else:
                    para.add_run(val)

print("✓ Table 5 (Comparison) updated")


# ─────────────────────────────────────────────
# 4. Update key paragraph sentences
# ─────────────────────────────────────────────
replacements = [
    # Training config text
    ("trained for 50 epochs on an NVIDIA Tesla P100 GPU (16 GB VRAM)",
     "trained in two phases: Phase 1 — 50 epochs on NVIDIA Tesla P100 (16 GB), "
     "Phase 2 — 80 additional epochs (resume) on 2× NVIDIA Tesla T4 (15 GB each), "
     "totalling 130 epochs"),
    ("AdamW optimiser was selected for its decoupled weight decay (0.0005)",
     "AdamW optimiser (Phase 1: lr0=0.0005, lrf=0.005; Phase 2: lr0=0.001, lrf=0.002) "
     "with decoupled weight decay (0.0005)"),
    # Results text
    ("overall mAP@0.5 of 0.419",   "overall mAP@0.5 of 0.616 (mAP@0.5:0.95 = 0.368)"),
    ("exceeding the research target of 0.4",
     "a significant improvement from Phase 1 mAP@0.5 = 0.427 (+44.3%)"),
    ("Car remains the strongest performer (mAP 0.52)",
     "Motorbike achieves the highest AP@0.5 (0.887) benefiting from the largest class count (16,312 instances). "
     "Car (mAP 0.776)"),
    # Comparison paragraph
    ("YOLOv11-L achieves a higher mAP (0.419 vs. 0.398 for YOLOv8-L)",
     "YOLOv11-L achieves a significantly higher mAP (0.616 vs. 0.398 for YOLOv8-L)"),
    # Conclusion paragraph
    ("mAP@0.5 of 0.419",  "mAP@0.5 of 0.616"),
]

for old, new in replacements:
    n = replace_in_all_paras(old, new)
    if n:
        print(f"✓ Replaced '{old[:50]}...' ({n} para(s))")
    else:
        print(f"⚠ Not found: '{old[:60]}'")


# ─────────────────────────────────────────────
# 5. Save
# ─────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"\n✅ Saved updated document to:\n   {OUTPUT_PATH}")
