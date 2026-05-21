# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'E:\test\DATN\YOLOv11_BaoCao_TiengViet_v3.docx')

def set_cell_text(cell, text):
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ''
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

def clear_and_set_para(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    for run in para.runs:
        run.text = ''
    para.runs[0].text = new_text

# ============================================================
# THAY DOI CAC DOAN VAN BAN (PARAGRAPHS)
# ============================================================
for i, para in enumerate(doc.paragraphs):
    text = para.text

    # Muc tieu (5) - so sanh YOLOv5/v8
    if 'YOLOv5-L và YOLOv8-L trên cùng bộ dữ liệu kiểm thử' in text:
        clear_and_set_para(para,
            '(5) So sánh phiên bản YOLOv11-L Cơ bản (huấn luyện trên FishEye8K, fine-tune toàn bộ) '
            'với phiên bản Nâng cao (bổ sung VisDrone2019, áp dụng SAHI và đóng băng 10 lớp backbone đầu), '
            'đánh giá đóng góp của từng kỹ thuật cải tiến.')

    # Dong gop chinh - mAP so sanh v5/v8
    elif 'vượt trội so với YOLOv5-L (0,361) và YOLOv8-L (0,398) trong cùng điều kiện' in text:
        clear_and_set_para(para,
            '- Phiên bản YOLOv11-L Cơ bản đạt mAP@0.5 = 0,419 trên tập kiểm thử FishEye8K. '
            'Phiên bản Nâng cao (bổ sung VisDrone-fisheye + SAHI + đóng băng 10 lớp backbone) '
            'đạt mAP@0.5 = 0,949 – cải thiện 126,5% so với phiên bản Cơ bản.')

    # Giai doan 5 - so sanh baseline
    elif 'So sánh kết quả với baseline YOLOv5-L và YOLOv8-L' in text:
        clear_and_set_para(para,
            'Giai đoạn 5 – Đánh giá và so sánh: So sánh phiên bản YOLOv11-L Cơ bản và Nâng cao, '
            'phân tích đóng góp của dữ liệu VisDrone-fisheye, kỹ thuật đóng băng backbone và SAHI.')

    # Section header 4.5
    elif text.strip() == '4.5. So sánh với các phiên bản YOLO khác':
        clear_and_set_para(para, '4.5. So sánh hai phiên bản YOLOv11-L')

    # TOC entry 4.5
    elif '4.5. So sánh với các phiên bản YOLO khác' in text and '....' in text:
        clear_and_set_para(para,
            text.replace('4.5. So sánh với các phiên bản YOLO khác',
                         '4.5. So sánh hai phiên bản YOLOv11-L'))

    # Intro section 4.5 - "tiến hành so sánh với hai baseline"
    elif 'tiến hành so sánh với hai baseline là YOLOv5-L và YOLOv8-L' in text:
        clear_and_set_para(para,
            'Để đánh giá đóng góp của từng kỹ thuật trong pipeline huấn luyện, đề tài so sánh hai '
            'phiên bản YOLOv11-L: Phiên bản Cơ bản chỉ sử dụng bộ dữ liệu FishEye8K với fine-tune '
            'toàn bộ mô hình (50 epoch, AdamW, img_size=640); Phiên bản Nâng cao bổ sung dữ liệu '
            'VisDrone2019 đã chuyển đổi fisheye, đóng băng 10 lớp backbone đầu, img_size=960 và áp '
            'dụng SAHI trong inference (80 epoch, SGD Cosine LR).')

    # Analysis after comparison table - "YOLOv11-L vượt trội rõ ràng"
    elif 'YOLOv11-L vượt trội rõ ràng về mAP@0.5 (+5,8% so với YOLOv8-L' in text:
        clear_and_set_para(para,
            'Phiên bản YOLOv11-L Nâng cao vượt trội rõ rệt so với phiên bản Cơ bản trên toàn bộ '
            'chỉ số đánh giá: mAP@0.5 tăng từ 0,419 lên 0,949 (+126,5%), Precision từ 0,65 lên '
            '0,931 (+43,2%), Recall từ 0,57 lên 0,899 (+57,7%). Kết quả khẳng định hiệu quả của '
            'sự kết hợp: bổ sung VisDrone-fisheye mở rộng đa dạng ngữ cảnh huấn luyện; đóng băng '
            'backbone bảo toàn đặc trưng tiền huấn luyện; SAHI cải thiện đáng kể khả năng phát '
            'hiện đối tượng nhỏ trong ảnh fisheye.')

    # Note about recall - "Recall của YOLOv11-L (0,57) thấp hơn YOLOv8-L"
    elif 'Recall của YOLOv11-L (0,57) thấp hơn YOLOv8-L (0,65)' in text:
        clear_and_set_para(para,
            'Đáng chú ý, Recall lớp Pedestrian trong phiên bản Cơ bản chỉ đạt 0,32 – thách thức '
            'lớn nhất của ảnh fisheye góc cao khi người đi bộ chiếm diện tích rất nhỏ. Phiên bản '
            'Nâng cao cải thiện Recall Pedestrian lên 0,815 (+154,7%) nhờ dữ liệu VisDrone phong '
            'phú hơn và SAHI. Tuy nhiên FPS giảm từ 41 xuống ~6 khi dùng SAHI, phù hợp cho phân '
            'tích offline hơn là realtime.')

    # Training time comparison
    elif 'YOLOv5-L cần ~3,2 giờ, YOLOv8-L cần ~4,1 giờ' in text:
        clear_and_set_para(para,
            'Thời gian huấn luyện: Phiên bản Cơ bản ~3,8 giờ (50 epoch, img_size=640, Tesla '
            'P100-16GB); Phiên bản Nâng cao ~6,8 giờ (80 epoch, img_size=960, Tesla P100-16GB). '
            'Inference với SAHI chậm hơn ~7× so với standard inference do chia lát ảnh.')

    # Ket luan - mAP so v5/v8
    elif 'vượt trội YOLOv5-L (0,361, +16,1%) và YOLOv8-L (0,398, +5,3%)' in text:
        clear_and_set_para(para,
            '• Fine-tune YOLOv11-L phiên bản Cơ bản (FishEye8K) đạt mAP@0.5 = 0,419; '
            'phiên bản Nâng cao (VisDrone + đóng băng backbone + SAHI) đạt mAP@0.5 = 0,949 '
            '– cải thiện 126,5%.')

    # Ket luan - params so v8
    elif '25,3M tham số và 86,9 GFLOPs – giảm 42% tham số và 47% GFLOPs so với YOLOv8-L' in text:
        clear_and_set_para(para,
            '• YOLOv11-L sử dụng 25,3M tham số và 86,9 GFLOPs. Kỹ thuật đóng băng backbone '
            '(10 lớp đầu) trong phiên bản Nâng cao giúp bảo toàn đặc trưng tiền huấn luyện '
            'và tránh overfitting khi mở rộng dataset, đồng thời tăng tốc độ hội tụ.')

    # Caption Bảng 4.7
    elif 'Bảng 4.7' in text and ('YOLOv5' in text or 'YOLOv8' in text):
        new_text = text.replace(
            'So sánh YOLOv5-L, YOLOv8-L và YOLOv11-L trên cùng bộ test',
            'So sánh YOLOv11-L Cơ bản và YOLOv11-L Nâng cao')
        clear_and_set_para(para, new_text)

# ============================================================
# TABLE 2 - Danh muc bang: cap nhat ten Bang 4.7
# ============================================================
tables = doc.tables
table2 = tables[2]
for row in table2.rows:
    if len(row.cells) >= 2 and 'So sánh YOLOv5-L' in row.cells[1].text:
        set_cell_text(row.cells[1], 'So sánh YOLOv11-L Cơ bản và YOLOv11-L Nâng cao')

# ============================================================
# TABLE 14 - Bang 4.5: Sieu tham so huan luyen (2 phien ban)
# ============================================================
table14 = tables[14]

# Doi ten cot header
set_cell_text(table14.rows[0].cells[0], 'Siêu tham số')
set_cell_text(table14.rows[0].cells[1], 'Phiên bản Cơ bản\n(FishEye8K)')
set_cell_text(table14.rows[0].cells[2], 'Phiên bản Nâng cao\n(+VisDrone+SAHI)')

# Gia tri cua phien ban Nang cao tuong ung voi tung tham so
enhanced_vals = {
    'model':        'yolo11l.pt',
    'epochs':       '80',
    'batch_size':   '16',
    'img_size':     '960',
    'optimizer':    'SGD (Cosine LR)',
    'lr0':          '0,01',
    'lrf':          '0,01',
    'weight_decay': '0,0005',
    'momentum':     '0,937',
    'warmup_epochs':'3',
    'patience':     '50',
    'save_period':  '10',
    'cache':        'disk',
    'amp':          'True',
    'close_mosaic': '–',
    'mosaic':       '0,8',
    'mixup':        '0,15',
    'copy_paste':   '–',
    'degrees':      '10,0',
    'translate':    '0,1',
    'scale':        '0,4',
    'erasing':      '–',
    'hsv_s':        '0,7',
    'hsv_v':        '0,4',
    'fliplr':       '0,5',
    'flipud':       '0,0',
    'shear':        '2,0',
    'perspective':  '0,0005',
    'freeze':       '10 lớp đầu',
}

for row in table14.rows[1:]:
    param = row.cells[0].text.strip()
    if param in enhanced_vals:
        set_cell_text(row.cells[2], enhanced_vals[param])
    else:
        # Xoa giai thich cu, de trong cho phu hop
        set_cell_text(row.cells[2], '–')

# ============================================================
# TABLE 15 - Bang 4.6: Ket qua huan luyen (2 phien ban)
# ============================================================
table15 = tables[15]

# Cap nhat header de the hien 2 phien ban
set_cell_text(table15.rows[0].cells[0], 'Lớp đối tượng')
set_cell_text(table15.rows[0].cells[1], 'Precision\n(CB / NC)')
set_cell_text(table15.rows[0].cells[2], 'Recall\n(CB / NC)')
set_cell_text(table15.rows[0].cells[3], 'mAP@0.5\n(CB / NC)')
set_cell_text(table15.rows[0].cells[4], 'F1-Score\n(CB / NC)')

# Du lieu ket qua ca 2 phien ban (CB = Co ban, NC = Nang cao)
enhanced_results = {
    'Car':        ('0,71 / 0,950', '0,68 / 0,950', '0,52 / 0,977', '0,69 / 0,950'),
    'Bus':        ('0,58 / 0,934', '0,52 / 0,941', '0,41 / 0,969', '0,55 / 0,937'),
    'Truck':      ('0,55 / 0,934', '0,74 / 0,881', '0,42 / 0,940', '0,63 / 0,907'),
    'Pedestrian': ('0,78 / 0,903', '0,32 / 0,815', '0,38 / 0,900', '0,45 / 0,857'),
    'Motorbike':  ('0,63 / 0,936', '0,57 / 0,908', '0,45 / 0,960', '0,60 / 0,922'),
    'ALL (mean)': ('0,65 / 0,931', '0,57 / 0,899', '0,419 / 0,949', '0,61 / 0,915'),
}

for row in table15.rows[1:]:
    cls = row.cells[0].text.strip()
    if cls in enhanced_results:
        p, r, m, f = enhanced_results[cls]
        set_cell_text(row.cells[1], p)
        set_cell_text(row.cells[2], r)
        set_cell_text(row.cells[3], m)
        set_cell_text(row.cells[4], f)

# ============================================================
# TABLE 16 - Bang 4.7: So sanh (doi v5/v8/v11 -> Co ban / Nang cao)
# ============================================================
table16 = tables[16]

# Cap nhat header
set_cell_text(table16.rows[0].cells[0], 'Phiên bản')

# Hang 1: YOLOv5-L → YOLOv11-L Co ban
r1 = table16.rows[1]
set_cell_text(r1.cells[0], 'YOLOv11-L Cơ bản\n(FishEye8K)')
set_cell_text(r1.cells[1], '0,419')
set_cell_text(r1.cells[2], '0,363')
set_cell_text(r1.cells[3], '0,65')
set_cell_text(r1.cells[4], '0,57')
set_cell_text(r1.cells[5], '25,3')
set_cell_text(r1.cells[6], '86,9')
set_cell_text(r1.cells[7], '41')

# Hang 2: YOLOv8-L → YOLOv11-L Nang cao
r2 = table16.rows[2]
set_cell_text(r2.cells[0], 'YOLOv11-L Nâng cao\n(+VisDrone+SAHI+Freeze)')
set_cell_text(r2.cells[1], '0,949')
set_cell_text(r2.cells[2], '0,705')
set_cell_text(r2.cells[3], '0,931')
set_cell_text(r2.cells[4], '0,899')
set_cell_text(r2.cells[5], '25,3')
set_cell_text(r2.cells[6], '86,9')
set_cell_text(r2.cells[7], '~6 (SAHI)')

# Hang 3: YOLOv11-L → xoa di (bo row thu 3)
if len(table16.rows) > 3:
    row3_elem = table16.rows[3]._element
    row3_elem.getparent().remove(row3_elem)

# ============================================================
# LUU FILE
# ============================================================
output_path = r'E:\test\DATN\YOLOv11_BaoCao_TiengViet_v3_modified.docx'
doc.save(output_path)
print(f'Da luu: {output_path}')
