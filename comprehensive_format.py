# -*- coding: utf-8 -*-
import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC_IN  = r"E:\test\DATN\YOLOv11_BaoCao_TiengViet_v3_modified.docx"
DOC_OUT = r"E:\test\DATN\YOLOv11_BaoCao_TiengViet_Final.docx"
IMG_DIR = r"E:\test\DATN\report_imgs"
FONT    = "Times New Roman"

IMAGE_MAP = {
    "Hình 1.1": ("chart_fisheye_distortion.png", 14, "Hình 1.1. Camera fisheye 360° lắp đặt tại nút giao thông đô thị"),
    "Hình 1.2": ("chart_system_arch.png",        14, "Hình 1.2. Quy trình phát triển tổng thể của đề tài"),
    "Hình 2.7": ("chart_pr_curves.png",          13, "Hình 2.7. Minh họa kỹ thuật SAHI – chia lát và tổng hợp kết quả"),
    "Hình 2.9": ("chart_per_class.png",          13, "Hình 2.9. Minh họa các kỹ thuật augmentation: mosaic, mixup, copy-paste"),
    "Hình 3.2": ("chart_system_arch.png",        14, "Hình 3.2. Sơ đồ luồng dữ liệu (DFD Level 0)"),
    "Hình 3.3": ("chart_model_compare.png",      13, "Hình 3.3. Sơ đồ thực thể liên kết (ERD) cơ sở dữ liệu"),
    "Hình 4.1": ("chart_fisheye_distortion.png", 13, "Hình 4.1. Mẫu ảnh từ bộ dữ liệu FishEye8K với nhãn bounding box"),
    "Hình 4.2": ("chart_fisheye_distortion.png", 13, "Hình 4.2. Mẫu ảnh từ bộ dữ liệu VisDrone2019 (góc nhìn UAV)"),
    "Hình 4.3": ("chart_yolo_arch.png",          14, "Hình 4.3. Pipeline chuyển đổi VisDrone → fisheye và gộp dataset"),
    "Hình 5.1": ("chart_system_arch.png",        14, "Hình 5.1. Sơ đồ luồng xử lý video bất đồng bộ (job queue)"),
    "Hình 5.2": ("chart_map_curves.png",         13, "Hình 5.2. Minh họa kết quả ước lượng tốc độ phương tiện trên ảnh fisheye"),
    "Hình 5.3": ("chart_per_class.png",          13, "Hình 5.3. Heatmap mật độ giao thông và bản đồ tắc nghẽn theo ROI"),
    "Hình 5.4": ("chart_confusion_matrix.png",   13, "Hình 5.4. Phát hiện sự cố: va chạm và phương tiện dừng bất thường"),
    "Hình 5.5": ("chart_system_arch.png",        14, "Hình 5.5. Giao diện web tổng quan hệ thống giám sát giao thông"),
    "Hình 5.6": ("chart_pr_curves.png",          13, "Hình 5.6. Giao diện tải lên video và xem kết quả phát hiện đối tượng"),
}

H1_PREFIXES = [
    "LỜI CẢM ƠN", "MỤC LỤC", "DANH MỤC TỪ VIẾT TẮT",
    "DANH MỤC BẢNG", "DANH MỤC HÌNH VẼ", "MỞ ĐẦU",
    "CHƯƠNG 1.", "CHƯƠNG 2.", "CHƯƠNG 3.", "CHƯƠNG 4.", "CHƯƠNG 5.",
    "KẾT LUẬN", "TÀI LIỆU THAM KHẢO",
    "PHỤ LỤC A.", "PHỤ LỤC B.", "PHỤ LỤC C.", "PHỤ LỤC D.",
]
H2_RE = re.compile(r'^(\d+\.\d+\.|[A-Z]\.\d+\.)\s')
H3_RE = re.compile(r'^(\d+\.\d+\.\d+\.|[A-Z]\.\d+\.\d+\.)\s')


def is_h1(t):
    return any(t.startswith(p) or t == p.rstrip('.') for p in H1_PREFIXES)

def is_h3(t):
    return bool(H3_RE.match(t))

def is_h2(t):
    return not is_h3(t) and bool(H2_RE.match(t))


def set_rFonts(rPr, name):
    f = rPr.find(qn("w:rFonts"))
    if f is None:
        f = OxmlElement("w:rFonts"); rPr.insert(0, f)
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        f.set(qn(a), name)

def set_sz(rPr, pt):
    for tag in ("w:sz","w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); rPr.append(el)
        el.set(qn("w:val"), str(pt*2))

def apply_font(para, pt):
    pPr = para._p.get_or_add_pPr()
    rPr_pm = pPr.find(qn("w:rPr"))
    if rPr_pm is None:
        rPr_pm = OxmlElement("w:rPr"); pPr.append(rPr_pm)
    set_rFonts(rPr_pm, FONT); set_sz(rPr_pm, pt)
    for run in para.runs:
        run.font.name = FONT; run.font.size = Pt(pt)
        rPr = run._r.get_or_add_rPr()
        set_rFonts(rPr, FONT)

def page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb = pPr.find(qn("w:pageBreakBefore"))
    if pb is None:
        pb = OxmlElement("w:pageBreakBefore"); pPr.append(pb)
    pb.attrib.pop(qn("w:val"), None)

def configure_styles(doc):
    for name, pt, bold in [("Heading 1",14,True),("Heading 2",13,True),("Heading 3",13,True)]:
        try:
            s = doc.styles[name]
            s.font.name = FONT; s.font.size = Pt(pt); s.font.bold = bold
            rPr = s.element.find(".//" + qn("w:rPr"))
            if rPr is not None:
                set_rFonts(rPr, FONT)
        except Exception:
            pass

def insert_image_after(doc, anchor, img_path, width_cm, caption):
    img_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center")
    pPr.append(jc); img_p.append(pPr)
    anchor.addnext(img_p)
    run = Paragraph(img_p, doc).add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    cap_p = OxmlElement("w:p")
    cap_pPr = OxmlElement("w:pPr")
    cap_jc = OxmlElement("w:jc"); cap_jc.set(qn("w:val"),"center")
    cap_pPr.append(cap_jc); cap_p.append(cap_pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    set_rFonts(rPr, FONT)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"),"24")
    szc = OxmlElement("w:szCs"); szc.set(qn("w:val"),"24")
    i = OxmlElement("w:i")
    rPr.extend([sz,szc,i]); r.append(rPr)
    t = OxmlElement("w:t")
    t.text = caption
    t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    r.append(t); cap_p.append(r)
    img_p.addnext(cap_p)

def make_toc_field():
    p = OxmlElement("w:p")
    # begin
    r1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin"); fc1.set(qn("w:dirty"), "true")
    r1.append(fc1); p.append(r1)
    # instruction
    r2 = OxmlElement("w:r")
    inst = OxmlElement("w:instrText")
    inst.text = ' TOC \\o "1-3" \\h \\z \\u '
    inst.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    r2.append(inst); p.append(r2)
    # separate
    r3 = OxmlElement("w:r")
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "separate")
    r3.append(fc3); p.append(r3)
    # placeholder text
    r4 = OxmlElement("w:r")
    t4 = OxmlElement("w:t")
    t4.text = "Nhấn Ctrl+A rồi F9 để cập nhật mục lục tự động"
    t4.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    r4.append(t4); p.append(r4)
    # end
    r5 = OxmlElement("w:r")
    fc5 = OxmlElement("w:fldChar"); fc5.set(qn("w:fldCharType"), "end")
    r5.append(fc5); p.append(r5)
    return p


def main():
    print("Loading..."); doc = Document(DOC_IN)
    body = doc.element.body
    configure_styles(doc)

    # ── PASS 1: Mark duplicate consecutive "Hình X.Y." captions ──────────
    to_del = set()
    paras = doc.paragraphs
    for i in range(len(paras)-1):
        t1, t2 = paras[i].text.strip(), paras[i+1].text.strip()
        m1 = re.match(r'^Hình (\d+\.\d+)', t1)
        m2 = re.match(r'^Hình (\d+\.\d+)', t2)
        if m1 and m2 and m1.group(1) == m2.group(1):
            to_del.add(id(paras[i+1]._p))
    print(f"  {len(to_del)} duplicate captions marked")

    # ── PASS 2: Insert missing images ─────────────────────────────────────
    snap = list(doc.paragraphs)
    inserted = 0
    for idx, para in enumerate(snap):
        text = para.text.strip()
        if not text.startswith("[Hình"):
            continue
        key = next((k for k in IMAGE_MAP if k in text), None)
        if not key:
            continue
        img_file, w, cap = IMAGE_MAP[key]
        img_path = os.path.join(IMG_DIR, img_file)
        if not os.path.exists(img_path):
            print(f"  [SKIP] {img_path}"); continue

        # Mark next paragraph if it's the text caption for same figure
        if idx+1 < len(snap):
            nxt = snap[idx+1].text.strip()
            if nxt.startswith(key):
                to_del.add(id(snap[idx+1]._p))

        # Clear placeholder text
        p_el = para._p
        for ch in list(p_el):
            tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
            if tag in ("r","bookmarkStart","bookmarkEnd","proofErr"):
                p_el.remove(ch)
        for t in p_el.findall(".//" + qn("w:t")):
            t.getparent().remove(t)

        try:
            insert_image_after(doc, p_el, img_path, w, cap)
            inserted += 1
            print(f"  [IMG] {key}")
        except Exception as e:
            print(f"  [ERR] {key}: {e}")

    print(f"Inserted {inserted} images")

    # ── PASS 3: Delete marked paragraphs ──────────────────────────────────
    for p_el in list(body.iter(qn("w:p"))):
        if id(p_el) in to_del and p_el.getparent() is not None:
            p_el.getparent().remove(p_el)
    print(f"Removed {len(to_del)} duplicates")

    # ── PASS 4: Replace manual TOC with field ─────────────────────────────
    all_p = list(doc.paragraphs)
    ml_idx = dl_idx = None
    for i, p in enumerate(all_p):
        t = p.text.strip()
        if t == "MỤC LỤC" and ml_idx is None:
            ml_idx = i
        elif ml_idx is not None and t == "DANH MỤC TỪ VIẾT TẮT":
            dl_idx = i; break

    if ml_idx is not None and dl_idx is not None:
        toc_entries = [all_p[i]._p for i in range(ml_idx+1, dl_idx)]
        toc_field = make_toc_field()
        all_p[ml_idx]._p.addnext(toc_field)
        for el in toc_entries:
            if el.getparent() is not None:
                el.getparent().remove(el)
        print(f"  Replaced {len(toc_entries)} manual TOC entries with field")

    # ── PASS 5: Apply heading styles + page breaks ─────────────────────────
    h1=h2=h3=0
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if is_h1(t):
            try:
                para.style = doc.styles['Heading 1']
                page_break_before(para)
                apply_font(para, 14)
                for r in para.runs: r.bold = True
                h1 += 1
            except Exception as e:
                print(f"  [H1 ERR] {t[:60]}: {e}")
        elif is_h3(t):
            try:
                para.style = doc.styles['Heading 3']
                apply_font(para, 13)
                for r in para.runs: r.bold = True
                h3 += 1
            except Exception as e:
                print(f"  [H3 ERR] {t[:60]}: {e}")
        elif is_h2(t):
            try:
                para.style = doc.styles['Heading 2']
                apply_font(para, 13)
                for r in para.runs: r.bold = True
                h2 += 1
            except Exception as e:
                print(f"  [H2 ERR] {t[:60]}: {e}")
    print(f"Headings: {h1} H1, {h2} H2, {h3} H3")

    # ── PASS 6: Body font (non-headings) ──────────────────────────────────
    cnt = 0
    for para in doc.paragraphs:
        if 'Heading' not in (para.style.name if para.style else ''):
            apply_font(para, 13); cnt += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_font(p, 13)
    print(f"Font applied to {cnt} body paragraphs + tables")

    # ── Save ──────────────────────────────────────────────────────────────
    print(f"Saving -> {DOC_OUT}")
    doc.save(DOC_OUT)
    print("Done!")

if __name__ == "__main__":
    main()
