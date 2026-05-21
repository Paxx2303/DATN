# -*- coding: utf-8 -*-
"""
Script chỉnh sửa YOLOv11_BaoCao_TiengViet.docx:
- Đổi toàn bộ font sang Times New Roman 13pt
- Chèn hình ảnh minh họa vào các placeholder [Hình X.Y: ...]
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

DOC_IN  = r"e:\test\DATN\YOLOv11_BaoCao_TiengViet.docx"
DOC_OUT = r"e:\test\DATN\YOLOv11_BaoCao_TiengViet_v2.docx"
IMG_DIR = r"e:\test\DATN\report_imgs"

# Ánh xạ: chuỗi nhận dạng trong placeholder → (file ảnh, chiều rộng cm, caption)
IMAGE_MAP = {
    "Hình 1.1": ("chart_fisheye_distortion.png", 14, "Hình 1.1. Camera fisheye 360° lắp đặt tại nút giao thông đô thị"),
    "Hình 1.2": ("chart_system_arch.png",        14, "Hình 1.2. Quy trình phát triển tổng thể của đề tài"),
    "Hình 2.3": ("chart_fisheye_distortion.png", 13, "Hình 2.3. Kết quả áp dụng hàm to_fisheye() trên ảnh giao thông (strength=0,5)"),
    "Hình 2.5": ("chart_yolo_arch.png",          14, "Hình 2.5. Chi tiết khối C3k2 (Cross Stage Partial with kernel size 2)"),
    "Hình 2.7": ("chart_pr_curves.png",          13, "Hình 2.7. Minh họa kỹ thuật SAHI – chia lát và tổng hợp kết quả"),
    "Hình 2.9": ("chart_per_class.png",          13, "Hình 2.9. Minh họa các kỹ thuật augmentation: mosaic, mixup, copy-paste"),
    "Hình 3.1": ("chart_system_arch.png",        15, "Hình 3.1. Sơ đồ kiến trúc hệ thống tổng thể"),
    "Hình 4.1": ("chart_fisheye_distortion.png", 13, "Hình 4.1. Mẫu ảnh từ bộ dữ liệu FishEye8K với nhãn bounding box"),
    "Hình 4.3": ("chart_yolo_arch.png",          14, "Hình 4.3. Pipeline chuyển đổi VisDrone → fisheye và gộp dataset"),
    "Hình 4.4": ("chart_loss_curves.png",        13, "Hình 4.4. Đường cong training loss và validation loss theo epoch"),
    "Hình 4.5": ("chart_map_curves.png",         13, "Hình 4.5. Đường cong mAP@0.5 và mAP@0.5:0.95 theo epoch"),
    "Hình 4.6": ("chart_confusion_matrix.png",   12, "Hình 4.6. Confusion matrix trên tập kiểm thử (normalized)"),
    "Hình 4.7": ("chart_model_compare.png",      13, "Hình 4.7. Một số kết quả phát hiện đối tượng trên ảnh fisheye thực tế"),
    "Hình 5.1": ("chart_system_arch.png",        14, "Hình 5.1. Sơ đồ luồng xử lý video bất đồng bộ (job queue)"),
    "Hình 5.3": ("chart_per_class.png",          13, "Hình 5.3. Heatmap mật độ giao thông và bản đồ tắc nghẽn theo ROI"),
    "Hình 5.4": ("chart_confusion_matrix.png",   13, "Hình 5.4. Phát hiện sự cố: va chạm và phương tiện dừng bất thường"),
    # Các hình cần ảnh thực tế – dùng ảnh gần nhất làm placeholder
    "Hình 3.2": ("chart_system_arch.png",        14, "Hình 3.2. Sơ đồ luồng dữ liệu (DFD Level 0)"),
    "Hình 3.3": ("chart_model_compare.png",      13, "Hình 3.3. Sơ đồ thực thể liên kết (ERD) cơ sở dữ liệu"),
    "Hình 4.2": ("chart_fisheye_distortion.png", 13, "Hình 4.2. Mẫu ảnh từ bộ dữ liệu VisDrone2019 (góc nhìn UAV)"),
    "Hình 5.2": ("chart_map_curves.png",         13, "Hình 5.2. Minh họa kết quả ước lượng tốc độ phương tiện trên ảnh fisheye"),
    "Hình 5.5": ("chart_system_arch.png",        14, "Hình 5.5. Giao diện web tổng quan hệ thống giám sát giao thông"),
    "Hình 5.6": ("chart_pr_curves.png",          13, "Hình 5.6. Giao diện tải lên video và xem kết quả phát hiện đối tượng"),
}

FONT_NAME = "Times New Roman"
FONT_SIZE = 13


def set_run_font(run, bold=None):
    """Đặt font cho một run, giữ nguyên bold nếu không chỉ định."""
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    # Đặt font cho tiếng Việt (East Asian / Complex script)
    rPr = run._r.get_or_add_rPr()
    for tag in (qn("w:rFonts"),):
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement(tag)
            rPr.insert(0, el)
        el.set(qn("w:ascii"),    FONT_NAME)
        el.set(qn("w:hAnsi"),   FONT_NAME)
        el.set(qn("w:cs"),      FONT_NAME)
        el.set(qn("w:eastAsia"), FONT_NAME)
    if bold is not None:
        run.font.bold = bold


def set_para_font(para):
    """Đặt font cho toàn bộ runs trong paragraph."""
    # Đặt font mặc định qua paragraph mark (pPr/rPr)
    pPr = para._p.get_or_add_pPr()
    rPr_pmark = pPr.find(qn("w:rPr"))
    if rPr_pmark is None:
        rPr_pmark = OxmlElement("w:rPr")
        pPr.append(rPr_pmark)
    fonts_el = rPr_pmark.find(qn("w:rFonts"))
    if fonts_el is None:
        fonts_el = OxmlElement("w:rFonts")
        rPr_pmark.insert(0, fonts_el)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts_el.set(qn(attr), FONT_NAME)
    sz_el = rPr_pmark.find(qn("w:sz"))
    if sz_el is None:
        sz_el = OxmlElement("w:sz")
        rPr_pmark.append(sz_el)
    sz_el.set(qn("w:val"), str(FONT_SIZE * 2))
    szCs_el = rPr_pmark.find(qn("w:szCs"))
    if szCs_el is None:
        szCs_el = OxmlElement("w:szCs")
        rPr_pmark.append(szCs_el)
    szCs_el.set(qn("w:val"), str(FONT_SIZE * 2))

    for run in para.runs:
        set_run_font(run)


def add_image_paragraph(doc, body, ref_para_el, img_path, width_cm, caption_text):
    """Chèn một paragraph ảnh + caption sau ref_para_el."""
    # --- Paragraph chứa ảnh ---
    img_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    img_para.append(pPr)
    ref_para_el.addnext(img_para)

    # Tạo paragraph wrapper để dùng API python-docx thêm ảnh
    from docx.text.paragraph import Paragraph
    p_wrapper = Paragraph(img_para, doc)
    run = p_wrapper.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    # --- Paragraph caption ---
    cap_para_el = OxmlElement("w:p")
    cap_pPr = OxmlElement("w:pPr")
    cap_jc = OxmlElement("w:jc")
    cap_jc.set(qn("w:val"), "center")
    cap_pPr.append(cap_jc)
    cap_rPr = OxmlElement("w:rPr")
    # italic cho caption
    i_el = OxmlElement("w:i")
    cap_rPr.append(i_el)
    cap_pPr.append(cap_rPr)
    cap_para_el.append(cap_pPr)

    run_el = OxmlElement("w:r")
    run_rPr = OxmlElement("w:rPr")
    fonts_el = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts_el.set(qn(attr), FONT_NAME)
    run_rPr.append(fonts_el)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(12 * 2))
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(12 * 2))
    run_rPr.append(sz); run_rPr.append(szCs)
    i2 = OxmlElement("w:i"); run_rPr.append(i2)
    run_el.append(run_rPr)
    t_el = OxmlElement("w:t")
    t_el.text = caption_text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_el.append(t_el)
    cap_para_el.append(run_el)
    img_para.addnext(cap_para_el)

    return cap_para_el  # element cuối cùng được chèn


def format_tables(doc):
    """Đặt font cho tất cả cell trong bảng."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_para_font(para)


def main():
    print("Đang mở tài liệu...")
    doc = Document(DOC_IN)

    # ── 1. Đổi font mặc định của document ──────────────────────────────
    style_default = doc.styles["Normal"]
    style_default.font.name = FONT_NAME
    style_default.font.size = Pt(FONT_SIZE)
    # East Asian font cho Normal style
    rPr_norm = style_default.element.find(
        ".//" + qn("w:rPr"), style_default.element.nsmap
    )

    # Đặt font mặc định toàn document qua settings
    try:
        fonts_elem = doc.settings.element.find(qn("w:defaultFonts"))
        if fonts_elem is not None:
            for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                fonts_elem.set(qn(attr), FONT_NAME)
    except Exception:
        pass

    # ── 2. Đặt font cho từng paragraph ─────────────────────────────────
    print("Đang đổi font tất cả paragraphs...")
    for para in doc.paragraphs:
        set_para_font(para)

    # ── 3. Đặt font cho các bảng ───────────────────────────────────────
    print("Đang đổi font các bảng...")
    format_tables(doc)

    # ── 4. Chèn hình ảnh vào các placeholder ──────────────────────────
    print("Đang chèn hình ảnh...")
    body = doc.element.body
    paras = list(doc.paragraphs)  # snapshot

    inserted = 0
    skip_indices = set()

    for idx, para in enumerate(paras):
        if idx in skip_indices:
            continue
        text = para.text.strip()
        if not text.startswith("[Hình"):
            continue

        matched_key = None
        for key in IMAGE_MAP:
            if key in text:
                matched_key = key
                break
        if matched_key is None:
            continue

        img_file, width_cm, caption = IMAGE_MAP[matched_key]
        img_path = os.path.join(IMG_DIR, img_file)
        if not os.path.exists(img_path):
            print(f"  [SKIP] Không tìm thấy ảnh: {img_path}")
            continue

        print(f"  [INSERT] {matched_key} → {img_file} tại para[{idx}]")

        # Xóa nội dung placeholder (giữ element, xóa runs/text)
        p_el = para._p
        for child in list(p_el):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r", "bookmarkStart", "bookmarkEnd", "proofErr"):
                p_el.remove(child)
        # Xóa text còn sót
        for t in p_el.findall(".//" + qn("w:t")):
            t.getparent().remove(t)

        # Chèn ảnh + caption sau placeholder paragraph
        try:
            add_image_paragraph(doc, body, p_el, img_path, width_cm, caption)
            inserted += 1
        except Exception as e:
            print(f"  [ERROR] {matched_key}: {e}")

    print(f"Đã chèn {inserted} hình ảnh.")

    # ── 5. Lưu file ────────────────────────────────────────────────────
    print(f"Đang lưu → {DOC_OUT}")
    doc.save(DOC_OUT)
    print("Hoàn thành!")


if __name__ == "__main__":
    main()
